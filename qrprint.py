from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ------------------------------------------------
# 기본 설정
# ------------------------------------------------
QR_DIR = "qr_codes"                # QR코드 이미지 폴더
OUTPUT_FILE = "QR_All_8x4_A4final.docx"

ROWS = 8                           # 행 수
COLUMNS = 4                        # 열 수
LABELS_PER_PAGE = ROWS * COLUMNS   # 페이지당 QR 개수

# ------------------------------------------------
# A4 용지 고정 (python-docx 기본은 Letter이므로 직접 지정)
# ------------------------------------------------
A4_WIDTH = Inches(8.27)            # 210mm
A4_HEIGHT = Inches(11.69)          # 297mm

# 여백 설정
TOP_MARGIN = 1.4 / 2.54            # 1.4cm
BOTTOM_MARGIN = 1.2 / 2.54         # 1.2cm
SIDE_MARGIN = 0.7 / 2.54           # 0.7cm

# ------------------------------------------------
# 셀 크기 및 스타일 설정 (A4 세로 기준 8×4 맞춤)
# ------------------------------------------------
CELL_WIDTH = Inches(2.03)          # 약 51.6mm
CELL_HEIGHT = Inches(1.307)        # 약 33.2mm
IMAGE_SIZE = Inches(0.87)          # QR코드 약 22mm
FONT_SIZE = 8                      # 코드번호 글자 크기

# ------------------------------------------------
# 셀 내부 여백 제거 함수
# ------------------------------------------------
def set_cell_padding(cell, padding=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'bottom', 'left', 'right'):
        node = tcPr.xpath(f'w:{side}')[0] if tcPr.xpath(f'w:{side}') else None
        if not node:
            from docx.oxml.shared import OxmlElement
            node = OxmlElement(f'w:{side}')
            tcPr.append(node)
        node.set('w', str(padding))
        node.set('type', 'dxa')

# ------------------------------------------------
# 셀에 QR 이미지 + 코드번호 삽입
# ------------------------------------------------
def fill_table_cells(table, qr_files):
    flat_cells = [cell for row in table.rows for cell in row.cells]
    for i, cell in enumerate(flat_cells):
        if i >= len(qr_files):
            continue
        qr_path = qr_files[i]
        set_cell_padding(cell, 0)
        cell.vertical_alignment = 1

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(0)
        p.space_after = Pt(0)

        if os.path.exists(qr_path):
            run = p.add_run()
            run.add_picture(qr_path, width=IMAGE_SIZE, height=IMAGE_SIZE)

        # 파일명에서 확장자 제거 후 코드번호 표시
        name = os.path.splitext(os.path.basename(qr_path))[0]
        p2 = cell.add_paragraph(name)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(FONT_SIZE)
        p2.space_before = Pt(0)
        p2.space_after = Pt(0)

# ------------------------------------------------
# QR 이미지 불러오기
# ------------------------------------------------
qr_files = sorted(
    [os.path.join(QR_DIR, f) for f in os.listdir(QR_DIR) if f.lower().endswith(".png")]
)
total = len(qr_files)
print(f"총 {total}개의 QR 이미지 발견")

# ------------------------------------------------
# 문서 생성 및 A4 페이지 설정
# ------------------------------------------------
doc = Document()
section = doc.sections[0]

# 페이지 크기 A4로 강제 지정
section.page_width = A4_WIDTH
section.page_height = A4_HEIGHT

# 여백 설정
section.top_margin = Inches(TOP_MARGIN)
section.bottom_margin = Inches(BOTTOM_MARGIN)
section.left_margin = Inches(SIDE_MARGIN)
section.right_margin = Inches(SIDE_MARGIN)

# ------------------------------------------------
# 페이지 단위로 QR라벨 생성
# ------------------------------------------------
for start in range(0, total, LABELS_PER_PAGE):
    page_files = qr_files[start:start + LABELS_PER_PAGE]
    if start > 0:
        doc.add_page_break()

    table = doc.add_table(rows=ROWS, cols=COLUMNS)
    table.autofit = False

    for row in table.rows:
        tr = row._tr
        tr.height = CELL_HEIGHT
        tr.height_rule = 1
        for cell in row.cells:
            cell.width = CELL_WIDTH
            set_cell_padding(cell, 0)

    fill_table_cells(table, page_files)
    print(f"{start//LABELS_PER_PAGE + 1}페이지 완료 ({len(page_files)}개 QR)")

# ------------------------------------------------
# 불필요한 빈 문단 제거 (빈 페이지 방지)
# ------------------------------------------------
for p in doc.paragraphs:
    if not p.text.strip():
        p._element.getparent().remove(p._element)

# ------------------------------------------------
# 파일 저장
# ------------------------------------------------
doc.save(OUTPUT_FILE)
print(f"\n✅ A4 8×4 완벽 정렬 + 빈 페이지 제거 완료 → {OUTPUT_FILE}")
