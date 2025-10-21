from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import os

QR_DIR = "qr_codes"
OUTPUT_FILE = "QR_All_8x4_A4final.docx"

ROWS = 8
COLUMNS = 4
LABELS_PER_PAGE = ROWS * COLUMNS

A4_WIDTH_IN = 8.27
A4_HEIGHT_IN = 11.69
A4_WIDTH = Inches(A4_WIDTH_IN)
A4_HEIGHT = Inches(A4_HEIGHT_IN)

TOP_MARGIN = 0.8 / 2.54
BOTTOM_MARGIN = 0.8 / 2.54
SIDE_MARGIN = 0.7 / 2.54

FONT_SIZE = 8
EMU_PER_INCH = 914400

def set_cell_padding(cell, padding=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ('top', 'bottom', 'left', 'right'):
        nodes = tcPr.xpath(f'w:{side}')
        node = nodes[0] if nodes else None
        if node is None:
            node = OxmlElement(f'w:{side}')
            tcPr.append(node)
        node.set('w', str(int(padding)))
        node.set('type', 'dxa')

def set_table_fixed_layout(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    existing = tblPr.xpath('w:tblLayout')
    if existing:
        tblLayout = existing[0]
    else:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

    existing_mar = tblPr.xpath('w:tblCellMar')
    if existing_mar:
        try:
            tblPr.remove(existing_mar[0])
        except Exception:
            pass
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('top','bottom','left','right'):
        node = OxmlElement(f'w:{side}')
        node.set('w', '0')
        node.set('type', 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)

def fill_table_cells(table, qr_files, image_size_length):
    flat_cells = [cell for row in table.rows for cell in row.cells]
    for i, cell in enumerate(flat_cells):
        if i >= len(qr_files):
            continue
        qr_path = qr_files[i]
        set_cell_padding(cell, 0)
        try:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            cell.vertical_alignment = 1

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        if os.path.exists(qr_path):
            run = p.add_run()
            run.add_picture(qr_path, width=image_size_length, height=image_size_length)

        name = os.path.splitext(os.path.basename(qr_path))[0]
        p2 = cell.add_paragraph(name)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        if p2.runs:
            p2.runs[0].font.size = Pt(FONT_SIZE)
        else:
            r = p2.add_run(name)
            r.font.size = Pt(FONT_SIZE)

if not os.path.isdir(QR_DIR):
    print(f"폴더가 없습니다: {QR_DIR}")
    qr_files = []
else:
    qr_files = sorted(
        [os.path.join(QR_DIR, f) for f in os.listdir(QR_DIR) if f.lower().endswith(".png")]
    )
total = len(qr_files)
print(f"총 {total}개의 QR 이미지 발견")

doc = Document()

# 빈 기본 문단(처음 생성되는 빈 문단)이 있으면 제거해서 테이블이 정확히 맨 위부터 시작하도록 함
for p in list(doc.paragraphs):
    if not p.text.strip():
        p._element.getparent().remove(p._element)

section = doc.sections[0]
section.page_width = A4_WIDTH
section.page_height = A4_HEIGHT

section.top_margin = Inches(TOP_MARGIN)
section.bottom_margin = Inches(BOTTOM_MARGIN)
section.left_margin = Inches(SIDE_MARGIN)
section.right_margin = Inches(SIDE_MARGIN)

section.header_distance = Inches(0)
section.footer_distance = Inches(0)

# 사용 가능 영역 EMU 계산
page_width_emu = int(A4_WIDTH_IN * EMU_PER_INCH)
page_height_emu = int(A4_HEIGHT_IN * EMU_PER_INCH)
left_margin_emu = int(SIDE_MARGIN * EMU_PER_INCH)
right_margin_emu = int(SIDE_MARGIN * EMU_PER_INCH)
top_margin_emu = int(TOP_MARGIN * EMU_PER_INCH)
bottom_margin_emu = int(BOTTOM_MARGIN * EMU_PER_INCH)

usable_width_emu = page_width_emu - left_margin_emu - right_margin_emu
usable_height_emu = page_height_emu - top_margin_emu - bottom_margin_emu

# 고정 셀 높이 3.4cm -> EMU (라벨지 규격에 맞게 고정)
fixed_cell_height_emu = int((3.4 / 2.54) * EMU_PER_INCH)

# 8행 * 3.4cm = 27.2cm이므로 강제로 고정 (라벨지 규격)
CELL_HEIGHT_EMU = fixed_cell_height_emu

CELL_WIDTH_EMU = usable_width_emu // COLUMNS

label_height_in = (FONT_SIZE + 2) / 72.0
padding_small_in = 0.01

cell_width_in = CELL_WIDTH_EMU / EMU_PER_INCH
cell_height_in = CELL_HEIGHT_EMU / EMU_PER_INCH

image_size_in = min(cell_width_in - padding_small_in, cell_height_in - label_height_in - padding_small_in)
if image_size_in < 0.06:
    image_size_in = 0.06
IMAGE_SIZE = Inches(image_size_in)

# 디버그: 계산값 로그
print(f"usable_height_in={usable_height_emu/EMU_PER_INCH:.3f} in")
print(f"셀 높이 고정 요청=3.4cm -> 적용 CELL_HEIGHT_in={cell_height_in*2.54:.2f} cm (in={cell_height_in:.4f})")
print(f"CELL_WIDTH_in={cell_width_in:.4f}, IMAGE_SIZE_in={image_size_in:.4f}")

for start in range(0, total, LABELS_PER_PAGE):
    page_files = qr_files[start:start + LABELS_PER_PAGE]
    if start > 0:
        doc.add_page_break()

    table = doc.add_table(rows=ROWS, cols=COLUMNS)
    table.autofit = False
    table.style = 'Table Grid'

    set_table_fixed_layout(table)

    # 열 너비/셀 너비/행 높이 정수(EMU)로 강제 설정
    for col in table.columns:
        for cell in col.cells:
            cell.width = int(CELL_WIDTH_EMU)

    for row in table.rows:
        row.height = int(CELL_HEIGHT_EMU)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in row.cells:
            cell.width = int(CELL_WIDTH_EMU)
            set_cell_padding(cell, 0)
            try:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            except Exception:
                cell.vertical_alignment = 1

    fill_table_cells(table, page_files, IMAGE_SIZE)
    print(f"{start//LABELS_PER_PAGE + 1}페이지 완료 ({len(page_files)}개 QR)")

for p in list(doc.paragraphs):
    if not p.text.strip():
        p._element.getparent().remove(p._element)

doc.save(OUTPUT_FILE)
print(f"\n✅ A4 8×4 완벽 정렬 + 빈 페이지 제거 완료 → {OUTPUT_FILE}")